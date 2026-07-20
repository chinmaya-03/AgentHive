"""
Document Parsing Utility for AgentHive.
Extracts clean plain text from PDF, DOCX, and TXT requirement files.
"""

import io
from pathlib import Path
from typing import Union

import fitz  # PyMuPDF
import docx  # python-docx

from app.core.exceptions import FileParsingError


def extract_text_from_pdf(file_input: Union[bytes, str, Path]) -> str:
    """
    Extracts text from a PDF file (path or bytes stream) using PyMuPDF (fitz).
    
    :param file_input: File path or raw bytes content
    :return: Cleaned extracted text string
    :raises FileParsingError: If PDF parsing fails or document is corrupt
    """
    try:
        if isinstance(file_input, (str, Path)):
            doc = fitz.open(str(file_input))
        elif isinstance(file_input, bytes):
            doc = fitz.open(stream=file_input, filetype="pdf")
        else:
            raise ValueError("Unsupported input type for PDF extraction.")

        text_blocks = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text and text.strip():
                text_blocks.append(text.strip())

        doc.close()
        
        extracted = "\n\n".join(text_blocks)
        if not extracted.strip():
            raise FileParsingError("PDF file contains no readable text content.")
            
        return extracted

    except FileParsingError:
        raise
    except Exception as e:
        raise FileParsingError(f"Failed to parse PDF document: {str(e)}")


def extract_text_from_docx(file_input: Union[bytes, str, Path]) -> str:
    """
    Extracts text from a DOCX file (path or bytes stream) using python-docx.
    
    :param file_input: File path or raw bytes content
    :return: Cleaned extracted text string
    :raises FileParsingError: If DOCX parsing fails or document is corrupt
    """
    try:
        if isinstance(file_input, (str, Path)):
            document = docx.Document(str(file_input))
        elif isinstance(file_input, bytes):
            document = docx.Document(io.BytesIO(file_input))
        else:
            raise ValueError("Unsupported input type for DOCX extraction.")

        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        
        # Also extract table text if present
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        extracted = "\n".join(paragraphs)
        if not extracted.strip():
            raise FileParsingError("DOCX file contains no readable text content.")

        return extracted

    except FileParsingError:
        raise
    except Exception as e:
        raise FileParsingError(f"Failed to parse DOCX document: {str(e)}")


def extract_text_from_txt(file_input: Union[bytes, str, Path]) -> str:
    """
    Extracts text from a plain TXT file (path or bytes stream).
    
    :param file_input: File path or raw bytes content
    :return: Cleaned extracted text string
    :raises FileParsingError: If decoding fails or file is empty
    """
    try:
        if isinstance(file_input, (str, Path)):
            path = Path(file_input)
            content = path.read_text(encoding="utf-8", errors="replace")
        elif isinstance(file_input, bytes):
            content = file_input.decode("utf-8", errors="replace")
        else:
            raise ValueError("Unsupported input type for TXT extraction.")

        if not content.strip():
            raise FileParsingError("TXT file is empty.")

        return content.strip()

    except FileParsingError:
        raise
    except Exception as e:
        raise FileParsingError(f"Failed to parse TXT document: {str(e)}")


def parse_document(file_name: str, file_content: Union[bytes, str, Path]) -> str:
    """
    Unified router for extracting text based on file extension (.pdf, .docx, .txt).
    
    :param file_name: Original file name with extension
    :param file_content: File path or raw bytes
    :return: Plain text
    """
    ext = Path(file_name).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_content)
    elif ext == ".docx":
        return extract_text_from_docx(file_content)
    elif ext in (".txt", ".md"):
        return extract_text_from_txt(file_content)
    else:
        raise FileParsingError(f"Unsupported file format '{ext}'. Supported formats: .pdf, .docx, .txt")
